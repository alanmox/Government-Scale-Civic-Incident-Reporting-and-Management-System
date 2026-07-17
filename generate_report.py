import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Arial'

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Cover Page
doc.add_paragraph('\n\n\n')
title = doc.add_paragraph('GOVERNMENT-SCALE CIVIC INCIDENT REPORTING AND MANAGEMENT SYSTEM (GCIRMS)')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(16)
    run.font.bold = True

doc.add_paragraph('\n\n')

cover_details = [
    ('Project Title:', 'Government-Scale Civic Incident Reporting and Management System (GCIRMS)'),
    ('Student Name:', '[Your Name Here]'),
    ('Registration Number:', '[Your Registration Number]'),
    ('Institution:', '[Your University]'),
    ('Programme:', 'B.Sc. Software Engineering / Computer Science'),
    ('Course:', '[Course Name]'),
    ('Lecturer:', '[Lecturer Name]'),
    ('Submission Date:', 'July 2026')
]

for label, value in cover_details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{label} ').bold = True
    p.add_run(value)

doc.add_page_break()

# 2. DECLARATION
add_heading(doc, '2. DECLARATION', 1)
doc.add_paragraph('I, [Your Name], declare that this project, titled "Government-Scale Civic Incident Reporting and Management System," is my original work and has not been submitted for any degree or examination at any other university. All sources of information and technologies used have been properly acknowledged.')

# 3. ACKNOWLEDGEMENT
add_heading(doc, '3. ACKNOWLEDGEMENT', 1)
doc.add_paragraph('I would like to express my sincere gratitude to my supervisor, [Lecturer Name], for their invaluable guidance and support throughout this project. I also extend my thanks to the faculty members of the Computer Science department for their academic instructions which provided the foundation for this enterprise-grade system.')

# 4. ABSTRACT
add_heading(doc, '4. ABSTRACT', 1)
doc.add_paragraph('The Government-Scale Civic Incident Reporting and Management System (GCIRMS) is a specialized web-based application designed to bridge the communication gap between citizens and government agencies. It enables citizens to report civic issues (e.g., broken infrastructure, water shortages, safety hazards) and provides government institutions with a robust, automated workflow engine to route, verify, assign, and resolve these incidents within strict Service Level Agreements (SLAs).')
doc.add_paragraph('Built entirely using pure PHP 8+ Object-Oriented Programming without external frameworks, the system demonstrates mastery of enterprise software architecture, manual MVC pattern implementation, strict security enforcement, and highly relational database design in MySQL.')

# 5. PROBLEM STATEMENT
add_heading(doc, '5. PROBLEM STATEMENT', 1)
doc.add_paragraph('Currently, citizens face significant hurdles when attempting to report civic issues to relevant authorities. The lack of a centralized digital reporting mechanism results in:')
probs = [
    'Misrouting of reports to the wrong departments.',
    'Lack of transparency for the citizen regarding the status of their report.',
    'Inefficient manual workflows leading to delayed resolution times and SLA breaches.',
    'Poor resource allocation by government agencies due to a lack of centralized analytical data.'
]
for prob in probs:
    doc.add_paragraph(prob, style='List Number')

# 6. OBJECTIVES
add_heading(doc, '6. OBJECTIVES', 1)
add_heading(doc, 'General Objective', 2)
doc.add_paragraph('To design, develop, and deploy a scalable, secure, and centralized civic incident management platform that automates the reporting, routing, and resolution workflow for government infrastructure issues.')

add_heading(doc, 'Specific Objectives', 2)
objs = [
    'Implement a Role-Based Access Control (RBAC) architecture to manage Citizens, Verification Officers, Agency Supervisors, and Administrators securely.',
    'Develop a strict State Machine Workflow Engine to guarantee that incident reports follow a mandatory lifecycle (Submitted -> Verified -> Assigned -> Resolved).',
    'Automate Agency Routing based on incident categories (e.g., Water issues route automatically to DAWASA, Road issues to TANROADS).',
    'Implement a Work Order Subsystem enabling officers to log progress, costs, and internal notes.',
    'Provide Real-time Analytics Dashboards customized for different administrative tiers.'
]
for obj in objs:
    doc.add_paragraph(obj, style='List Number')

# 7. SYSTEM ARCHITECTURE & METHODOLOGY
add_heading(doc, '7. SYSTEM ARCHITECTURE & METHODOLOGY', 1)
add_heading(doc, '7.1 Architecture Design', 2)
doc.add_paragraph('The system abandons commercial frameworks (like Laravel) to demonstrate core software engineering competence. It utilizes a Manual MVC (Model-View-Controller) architecture implemented from scratch:')
arch_points = [
    'Front Controller (index.php): All requests are routed through a single entry point.',
    'Router & Middleware Pipeline: Custom routing engine parsing URLs and executing security middleware (CSRF, Auth) before reaching controllers.',
    'Repository Pattern: Database operations are strictly abstracted away from business logic. Models never query the database directly.',
    'Service Layer: Complex business rules (Workflows, SLAs, Routing, File Uploads) are handled by dedicated Service classes.'
]
for point in arch_points:
    doc.add_paragraph(point, style='List Bullet')

add_heading(doc, '7.2 Security by Design', 2)
sec_points = [
    'Authentication: Argon2ID hashing, session regeneration, and AES-256-GCM encryption for sensitive internal files.',
    'Database: PDO Prepared Statements exclusively used to eliminate SQL Injection.',
    'Uploads: Strict MIME-type binary validation (finfo_file) to prevent malicious executable uploads. Files are stored outside the public web root and served via secure proxy.'
]
for point in sec_points:
    doc.add_paragraph(point, style='List Bullet')

# 8. CORE MODULES IMPLEMENTED
add_heading(doc, '8. CORE MODULES IMPLEMENTED', 1)
mods = [
    'Infrastructure & Security: Custom routing, middleware chains, and session management.',
    'User & Organization Management: Multi-tier hierarchy (Agencies, Departments, Roles).',
    'Incident Reporting: Multi-category reporting with SLA tracking and geolocation data.',
    'Workflow & Routing Engine: Immutable database logging of all state transitions and automated assignment logic.',
    'Officer Work Orders: Granular tracking of task completion percentage, internal communications, and cost estimation.',
    'REST API: JWT-ready Bearer token endpoints for external mobile application integration.'
]
for mod in mods:
    doc.add_paragraph(mod, style='List Number')

# 9. CONCLUSION
add_heading(doc, '9. CONCLUSION', 1)
doc.add_paragraph('The GCIRMS project successfully fulfills the requirements of a modern, enterprise-grade government platform. By developing the MVC framework from scratch, the project demonstrates a profound understanding of HTTP lifecycles, Object-Oriented design patterns (Singleton, Factory, Repository, Strategy), and advanced relational database normalization. The platform is scalable, secure, and ready for deployment on cloud infrastructure (AWS EC2).')

# 10. FUTURE ENHANCEMENTS
add_heading(doc, '10. FUTURE ENHANCEMENTS', 1)
fut = [
    'Implementation of SMS Notifications via local telecom gateways.',
    'Full integration of interactive Leaflet.js maps for spatial density clustering of incidents.',
    'Machine Learning layer to predict SLA breach probabilities based on historical agency performance data.'
]
for f in fut:
    doc.add_paragraph(f, style='List Bullet')

doc.save('GCIRMS_University_Report.docx')
print('Report generated successfully as GCIRMS_University_Report.docx')
